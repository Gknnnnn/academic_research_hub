#!/usr/bin/env python3
"""
make_jimf_reference_doc.py — Create pandoc reference DOCX with JIMF table styles
2026-04-09

Generates reference.docx for use with:
  pandoc ... --reference-doc=reference.docx -o output.docx

JIMF style guide:
- Times New Roman 12pt body
- Tables: no vertical lines, bold header row, 10pt font
- Table title: bold, above table
- Table notes: 9pt italic, below table
- Margins: 2.54 cm all sides (standard A4)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup: A4, 2.54 cm margins ──────────────────────────────────────────
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)

# ── Default paragraph style ───────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(12)

# ── Heading styles ────────────────────────────────────────────────────────────
for i in range(1, 4):
    h = doc.styles[f"Heading {i}"]
    h.font.name = "Times New Roman"
    h.font.bold = True
    h.font.size = Pt(14 - (i - 1))
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)

# ── Table style ───────────────────────────────────────────────────────────────
# Use 'Table Grid' as base and modify
if "Table Grid" in doc.styles:
    ts = doc.styles["Table Grid"]
    ts.font.name = "Times New Roman"
    ts.font.size = Pt(10)

# ── Sample content demonstrating styles ───────────────────────────────────────
doc.add_heading("Title", level=1)
p = doc.add_paragraph(
    "This reference document sets the pandoc style for JIMF submissions. "
    "Body text: Times New Roman 12pt, justified."
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading("Section Heading", level=2)
doc.add_paragraph("Normal body paragraph text.")

# ── Demo table with JIMF formatting ──────────────────────────────────────────
doc.add_paragraph("Table N. Title — Subtitle", style="Normal").runs[0].bold = True

tbl = doc.add_table(rows=4, cols=4)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr = tbl.rows[0]
for j, cell in enumerate(hdr.cells):
    cell.text = f"Header {j+1}"
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data rows
for i in range(1, 4):
    for j, cell in enumerate(tbl.rows[i].cells):
        cell.text = f"Cell {i},{j+1}"
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Remove vertical borders (JIMF: no vertical lines)
def remove_cell_borders(table, borders_to_remove=("left", "right", "insideV")):
    tbl_el = table._tbl
    for row in tbl_el.iter():
        if row.tag.endswith("}tr"):
            for tc in row.iter():
                if tc.tag.endswith("}tc"):
                    tcPr = tc.find(qn("w:tcPr"))
                    if tcPr is None:
                        tcPr = OxmlElement("w:tcPr")
                        tc.insert(0, tcPr)
                    tcBorders = tcPr.find(qn("w:tcBorders"))
                    if tcBorders is None:
                        tcBorders = OxmlElement("w:tcBorders")
                        tcPr.append(tcBorders)
                    for border_name in borders_to_remove:
                        border = tcBorders.find(qn(f"w:{border_name}"))
                        if border is None:
                            border = OxmlElement(f"w:{border_name}")
                            tcBorders.append(border)
                        border.set(qn("w:val"), "nil")

remove_cell_borders(tbl)

# Notes paragraph
notes = doc.add_paragraph("Notes: 9pt italic notes below table. *** p<0.01, ** p<0.05, * p<0.10.")
notes.runs[0].font.size = Pt(9)
notes.runs[0].font.italic = True
notes.runs[0].font.name = "Times New Roman"

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/Users/mehmetgokhanozdemir/Library/CloudStorage/OneDrive-Kişisel/Akademik_Arastirma/200-Manuscripts/210-Active/2026-Digital-Assets-Monetary-Substitution-EM/04-Manuscript/reference.docx"
doc.save(out)
print(f"Saved: {out}")
